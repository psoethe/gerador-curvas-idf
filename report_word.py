"""
Word (.docx) report generator for IDF Curve Calculator.
Produces a professional Memorial de Cálculo / Calculation Report
using python-docx, mirroring the structure of the PDF report.
"""
import io
import logging
import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from i18n import t

logger = logging.getLogger("viktor")

# ── Brand colours ─────────────────────────────────────────────────────────────
_BLUE_DARK  = RGBColor(0x1a, 0x52, 0x76)
_BLUE_MID   = RGBColor(0x29, 0x80, 0xb9)
_BLUE_LIGHT = RGBColor(0xd6, 0xea, 0xf8)
_WHITE      = RGBColor(0xff, 0xff, 0xff)
_GREY_DARK  = RGBColor(0x33, 0x41, 0x55)
_GREEN      = RGBColor(0x16, 0xa3, 0x4a)


# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    tcPr.append(shd)


# ── Formatting helpers ────────────────────────────────────────────────────────

def _section_header(doc: Document, text: str) -> None:
    """Blue-bar section header."""
    p   = doc.add_paragraph()
    run = p.add_run(f'  {text}')
    run.bold = True
    run.font.size  = Pt(11)
    run.font.color.rgb = _WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '1a5276')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)


def _para(doc: Document, text: str, bold: bool = False, italic: bool = False,
          size: int = 10, color: RGBColor = None,
          align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a styled paragraph."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    return p


def _add_df_table(doc: Document, df,
                  header_hex: str = '1a5276',
                  alt_hex: str    = 'd6eaf8') -> None:
    """Render a pandas DataFrame as a styled Word table."""
    rows, cols = df.shape
    tbl = doc.add_table(rows=rows + 1, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Header
    for j, col_name in enumerate(df.columns):
        cell = tbl.rows[0].cells[j]
        cell.text = str(col_name)
        r = cell.paragraphs[0].runs[0]
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = _WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, header_hex)

    # Data rows
    for i, (_, row_data) in enumerate(df.iterrows()):
        bg = alt_hex if i % 2 == 0 else 'ffffff'
        for j, val in enumerate(row_data):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = str(val)
            r = cell.paragraphs[0].runs[0]
            r.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(cell, bg)

    doc.add_paragraph()


def _add_figure(doc: Document, png_bytes: bytes, caption: str,
                width_cm: float = 15.0) -> None:
    """Embed a PNG figure with a centred caption."""
    if not png_bytes:
        return
    buf = io.BytesIO(png_bytes)
    doc.add_picture(buf, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = _GREY_DARK
    doc.add_paragraph()


def _sherman_block(doc: Document, sherman: dict, lang: str) -> None:
    """
    Render the Sherman equation block:
      i = A · TR^B / (t ± C)^D
    with parameter grid and R² / RMSE / NSE.
    """
    import pandas as pd
    is_pt = lang == 'PT'
    A = sherman['A'];  B = sherman['B']
    C = sherman['C'];  D = sherman['D']
    r2   = sherman['R²']
    rmse = sherman['RMSE']
    nse  = sherman['NSE']
    c_sign = '+' if C >= 0 else '-'
    c_abs  = abs(C)

    # Equation line with superscripts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

    def _r(text, bold=True, size=13, color=_BLUE_DARK, sup=False):
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        if sup:
            run.font.superscript = True
        return run

    _r('i  =  ')
    _r(f'{A:.4f}')
    _r(' \u00b7 TR')          # · TR
    _r(f'{B:.4f}', sup=True, size=9)
    _r('  /  (t ')
    _r(f'{c_sign} {c_abs:.4f}')
    _r(')')
    _r(f'{D:.4f}', sup=True, size=9)

    # Parameter grid (1 row labels + 1 row values, 4 columns)
    lbl_a = 'A (CONSTANTE)'      if is_pt else 'A (CONSTANT)'
    lbl_b = 'B (EXPOENTE TR)'    if is_pt else 'B (TR EXPONENT)'
    lbl_c = 'C (AJUSTE TEMPO)'   if is_pt else 'C (TIME ADJUST)'
    lbl_d = 'D (EXPOENTE TEMPO)' if is_pt else 'D (TIME EXPONENT)'

    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    params_data = [(lbl_a, f'{A:.4f}'), (lbl_b, f'{B:.4f}'),
                   (lbl_c, f'{c_abs:.4f}'), (lbl_d, f'{D:.4f}')]
    for ci, (lbl, val) in enumerate(params_data):
        lc = tbl.rows[0].cells[ci]
        lc.text = lbl
        lc.paragraphs[0].runs[0].bold = True
        lc.paragraphs[0].runs[0].font.size = Pt(8)
        lc.paragraphs[0].runs[0].font.color.rgb = _WHITE
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(lc, '1e3a5f')
        vc = tbl.rows[1].cells[ci]
        vc.text = val
        vc.paragraphs[0].runs[0].bold = True
        vc.paragraphs[0].runs[0].font.size = Pt(11)
        vc.paragraphs[0].runs[0].font.color.rgb = _BLUE_DARK
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(vc, 'd6eaf8')
    doc.add_paragraph()

    # Validation metrics
    lbl_valid = 'Validacao Estatistica' if is_pt else 'Statistical Validation'
    _section_header(doc, f'📈 {lbl_valid}')

    r2_pct = f'{r2 * 100:.2f}%'
    lbl_r2 = 'Coeficiente de Determinacao (R2)' if is_pt else 'Coefficient of Determination (R2)'
    p_r2 = doc.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2_run = p_r2.add_run(f'{lbl_r2}:  {r2_pct}')
    r2_run.bold = True
    r2_run.font.size = Pt(14)
    r2_run.font.color.rgb = _GREEN

    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.style = 'Table Grid'
    for ci, (lbl, val) in enumerate([('RMSE (mm/h)', f'{rmse:.4f}'), ('NSE', f'{nse:.4f}')]):
        hc = tbl2.rows[0].cells[ci]
        hc.text = lbl
        hc.paragraphs[0].runs[0].bold = True
        hc.paragraphs[0].runs[0].font.size = Pt(9)
        hc.paragraphs[0].runs[0].font.color.rgb = _WHITE
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(hc, '1a5276')
        vc = tbl2.rows[1].cells[ci]
        vc.text = val
        vc.paragraphs[0].runs[0].bold = True
        vc.paragraphs[0].runs[0].font.size = Pt(11)
        vc.paragraphs[0].runs[0].font.color.rgb = _BLUE_DARK
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(vc, 'eaf4fb')
    doc.add_paragraph()


# ── Main public function ──────────────────────────────────────────────────────

def generate_word_report(
    results: dict,
    responsavel: str,
    localizacao: str,
    estacao: str,
    fig_historica,
    fig_gumbel,
    fig_pdf,
    fig_idf,
    lang: str = 'PT',
) -> bytes:
    """
    Generate a complete Word (.docx) Memorial de Calculo / Calculation Report.
    Returns raw .docx bytes.
    """
    import pandas as pd
    logger.info(f"📝 Gerando relatorio Word completo (lang={lang})...")

    # Unpack
    series_df  = results['series_df']
    gumbel_df  = results['gumbel_df']
    disagg_df  = results['disagg_df']
    idf_df     = results['idf_df']
    sherman    = results['sherman_params']
    mu         = results['mu']
    sigma      = results['sigma']
    n          = results['n_samples']
    isozona    = results['isozona']
    gumbel_mem = results['gumbel_memory']
    year_start = results.get('year_start')
    year_end   = results.get('year_end')

    from calculations import compute_taborga_memory, DURATIONS
    tab_mem = compute_taborga_memory(gumbel_mem, isozona, lang)

    col_ano = t('col_ano', lang)
    y0 = year_start or int(series_df[col_ano].min())
    y1 = year_end   or int(series_df[col_ano].max())
    period_str = f'{y0} - {y1}'

    is_pt = lang == 'PT'
    now   = datetime.datetime.now()
    report_id = f'IDF-{estacao or "XXX"}-{now.strftime("%Y%m%d-%H%M")}'

    # Render figures to PNG
    logger.info("🖼️ Renderizando figuras para PNG (Word)...")
    try:
        import plotly.io as pio
        png_hist   = pio.to_image(fig_historica, format='png', width=860, height=380, scale=1.5)
        png_gumbel = pio.to_image(fig_gumbel,    format='png', width=860, height=380, scale=1.5)
        png_pdf    = pio.to_image(fig_pdf,        format='png', width=860, height=440, scale=1.5)
        png_idf    = pio.to_image(fig_idf,        format='png', width=860, height=440, scale=1.5)
        logger.info("🖼️ Figuras Word: todas renderizadas")
    except Exception as exc:
        logger.warning(f"⚠️ Falha ao renderizar figuras para Word: {exc}")
        png_hist = png_gumbel = png_pdf = png_idf = None

    # Create document
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── COVER ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run('IDF Curve Calculator')
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = _BLUE_DARK

    sub_text = 'Memorial de Calculo - Curvas IDF' if is_pt else 'Calculation Report - IDF Curves'
    sp = doc.add_paragraph(sub_text)
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.runs[0].font.size = Pt(14)
    sp.runs[0].font.color.rgb = _BLUE_MID
    doc.add_paragraph()

    meta_rows = [
        (t('report_responsavel', lang), responsavel or '-'),
        (t('report_localizacao', lang),  localizacao or '-'),
        (t('report_estacao', lang),      estacao or '-'),
        (t('report_period', lang),       period_str),
        (t('report_isozona', lang),      isozona),
        (t('report_n', lang),            f'N = {n} {t("report_n_suffix", lang)}'),
        (t('report_date', lang),         now.strftime('%d/%m/%Y %H:%M')),
        ('ID', report_id),
    ]
    mt = doc.add_table(rows=len(meta_rows), cols=2)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    mt.style = 'Table Grid'
    for i, (k, v) in enumerate(meta_rows):
        bg = 'eaf4fb' if i % 2 == 0 else 'ffffff'
        kc = mt.rows[i].cells[0]
        vc = mt.rows[i].cells[1]
        kc.text = k
        vc.text = str(v)
        kc.paragraphs[0].runs[0].bold = True
        kc.paragraphs[0].runs[0].font.size = Pt(10)
        kc.paragraphs[0].runs[0].font.color.rgb = _BLUE_DARK
        vc.paragraphs[0].runs[0].font.size = Pt(10)
        _set_cell_bg(kc, 'd6eaf8')
        _set_cell_bg(vc, bg)

    doc.add_page_break()

    # ── PASSO 1 — Serie Historica ─────────────────────────────────────────────
    _section_header(doc, ('Passo 1 - Serie Historica de Precipitacao'
                          if is_pt else 'Step 1 - Historical Precipitation Series'))
    _para(doc, f'Estacao: {estacao or "-"}  |  Periodo: {period_str}  |  N = {n}')
    _para(doc, f'mu = {mu:.3f} mm  |  sigma = {sigma:.3f} mm', bold=True, color=_BLUE_DARK)
    doc.add_paragraph()
    _add_df_table(doc, series_df)
    _add_figure(doc, png_hist, 'Figura 1 - Serie Historica de Precipitacao Maxima Diaria Anual')
    doc.add_page_break()

    # ── PASSO 2 — Analise de Gumbel ───────────────────────────────────────────
    _section_header(doc, ('Passo 2 - Analise de Gumbel'
                          if is_pt else 'Step 2 - Gumbel Analysis'))
    _para(doc, f'N = {n}  |  Yn = {gumbel_mem["stats"]["yn"]:.4f}  |  Sn = {gumbel_mem["stats"]["sn"]:.4f}', bold=True)
    _para(doc, f'mu = {mu:.3f} mm  |  sigma = {sigma:.3f} mm')
    doc.add_paragraph()
    _add_df_table(doc, gumbel_df)
    _add_figure(doc, png_gumbel, 'Figura 2 - Analise de Gumbel')
    doc.add_page_break()

    # ── PASSO 3 — Desagregacao de Taborga ────────────────────────────────────
    _section_header(doc, ('Passo 3 - Desagregacao de Taborga'
                          if is_pt else 'Step 3 - Taborga Disaggregation'))
    _para(doc, f'Isozona: {isozona}', bold=True, color=_BLUE_DARK)
    doc.add_paragraph()
    k_rows = tab_mem['k_rows']
    if k_rows:
        k_df = pd.DataFrame(k_rows)
        k_df.columns = [
            t('pdf_k_tr', lang), t('pdf_k_p6', lang), t('pdf_k_p60', lang),
            t('pdf_k_p1440', lang), t('pdf_k_k1', lang), t('pdf_k_k2', lang),
        ]
        _add_df_table(doc, k_df)
    doc.add_page_break()

    # ── PASSO 4 — Curvas PDF ──────────────────────────────────────────────────
    _section_header(doc, ('Passo 4 - Curvas PDF - Precipitacao-Duracao-Frequencia'
                          if is_pt else 'Step 4 - PDF Curves - Precipitation-Duration-Frequency'))
    _para(doc, ('Precipitacao acumulada P(t) por duracao e periodo de retorno.'
                if is_pt else 'Accumulated precipitation P(t) by duration and return period.'))
    doc.add_paragraph()
    pdf_disp = disagg_df.copy()
    pdf_disp.index = DURATIONS
    pdf_disp.index.name = t('pdf_table_dur', lang)
    pdf_disp = pdf_disp.reset_index()
    pdf_disp.columns = [t('pdf_table_dur', lang)] + [
        f'TR={c.split("=")[1].split(" ")[0]}' for c in disagg_df.columns
    ]
    for c in pdf_disp.columns[1:]:
        pdf_disp[c] = pdf_disp[c].round(3)
    _add_df_table(doc, pdf_disp)
    _add_figure(doc, png_pdf, 'Figura 3 - Curvas PDF - Precipitacao Acumulada por Duracao')
    doc.add_page_break()

    # ── PASSO 5 — Curvas IDF ──────────────────────────────────────────────────
    _section_header(doc, ('Passo 5 - Curvas IDF - Intensidade-Duracao-Frequencia'
                          if is_pt else 'Step 5 - IDF Curves - Intensity-Duration-Frequency'))
    _para(doc, ('Intensidade i(t,TR) = P(t,TR) / (t/60) em mm/h.'
                if is_pt else 'Intensity i(t,TR) = P(t,TR) / (t/60) in mm/h.'))
    doc.add_paragraph()
    idf_disp = idf_df.copy()
    idf_disp.index = DURATIONS
    idf_disp.index.name = t('pdf_table_dur', lang)
    idf_disp = idf_disp.reset_index()
    idf_disp.columns = [t('pdf_table_dur', lang)] + [
        f'TR={c.split("=")[1].split(" ")[0]}' for c in idf_df.columns
    ]
    for c in idf_disp.columns[1:]:
        idf_disp[c] = idf_disp[c].round(3)
    _add_df_table(doc, idf_disp)
    _add_figure(doc, png_idf, 'Figura 4 - Curvas IDF - Intensidade-Duracao-Frequencia')
    doc.add_page_break()

    # ── PASSO 6 — Parametros de Sherman ──────────────────────────────────────
    _section_header(doc, ('Passo 6 - Parametros de Sherman'
                          if is_pt else 'Step 6 - Sherman Parameters'))
    _para(doc, ('Ajuste da equacao de Sherman (Montana) aos dados IDF por minimos quadrados nao-lineares.'
                if is_pt else 'Fitting the Sherman (Montana) equation to IDF data via nonlinear least squares.'))
    doc.add_paragraph()
    _sherman_block(doc, sherman, lang)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    _para(doc,
          (f'Relatorio gerado automaticamente pelo IDF Curve Calculator em {now.strftime("%d/%m/%Y %H:%M")}.'
           if is_pt else
           f'Report automatically generated by IDF Curve Calculator on {now.strftime("%d/%m/%Y %H:%M")}.'),
          italic=True, size=9, color=_GREY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f'ID: {report_id}', size=9, color=_GREY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Serialise
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    word_bytes = buf.read()
    logger.info(f"✅ Relatorio Word gerado: {len(word_bytes):,} bytes")
    return word_bytes
